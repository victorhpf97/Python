from docx import Document

document = Document()
document.add_heading('O que é:')
document.add_paragraph('     É um jogo eletrônico com o objetivo de testar o raciocínio e a agilidade do jogador , onde o mesmo deve encaixar blocos para preencher os espaços')

document.add_heading('Definição de níveis e Subníveis:')
document.add_paragraph('O game é composto por 16 níveis onde o jogador pode ir avançando de acordo com a sua pontuação')

document.add_heading('Requisitos Mínimos:')
document.add_paragraph('Sistema operacional Linux,Windows 500 MB de RAM , Placa de vídeo 2 Gb')

document.add_heading('Esquema de licenciamento:')
document.add_paragraph('Licença MIT open source sem fins lucrativos')

document.add_heading('Esquema de Atualização:')
document.add_paragraph('Versão 1.0.0 beta , atualização é feita mensalmente , favor conferir documento atualização.txt no diretório do jogo')

document.add_heading('Plano de manutenção:')
document.add_paragraph('A manutenção é feito pelas atualizações oficiais de acordo com a licença do usuário podendo o mesmo tirar duvidas sobre o produto através do site WWW.victorarts.com.br')

document.add_heading('Público alvo:')
document.add_paragraph('Destinado a todo tipo de publico e sem restrição de idade')

document.add_heading('Instalação:')
document.add_paragraph('abra o executável e leia as instruções e se estiver conforme pressione para avançar , um atalho será criado em sua área de trabalho')

document.add_heading('Informações sobre o autor:')
document.add_paragraph('VictorArts empresa desde de 1997 com sede nos estados unidos , desenvolvendo jogos eletrônicos para computadores e consoles.')

document.add_heading('Suporte:')
document.add_paragraph('para mais informações sobre o suporte técnico envie um e-mail para victorjogos@hotmail.com')

document.add_heading('Informações sobre a versões:')
document.add_paragraph('versão 1.0.0 beta ')

document.add_heading('Ajuda:')
document.add_paragraph('teclas : seta direcionais esquerda, direita, cima , baixo para movimentação das peças.')






document.save('D:/test.docx')